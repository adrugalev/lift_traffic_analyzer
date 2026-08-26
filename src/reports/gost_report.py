"""Отчёт по форме раздела 9 ГОСТ 34758-2021."""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from src import __version__
from src.models.project import Project
from src.models.results import CalculationResult, ComplianceStatus
from src.reports.docx_report import (
    COMPLIES_COLOR,
    DOES_NOT_COMPLY_COLOR,
    _add_list_item,
    _add_table,
    _configure_document,
    _set_document_language,
)
from src.reports.formula_rendering import (
    add_word_formula_block,
    pdf_formula_markup,
    pdf_formula_legend_markup,
    pdf_substitution_markup,
)
from src.reports.pdf_report import _register_fonts


GOST_BASIS = "GOST_34758_2021_CLAUSE_7"


def _format_report_datetime(value: datetime) -> str:
    """Форматирует локальную дату без служебного обозначения часового пояса."""

    return value.astimezone().strftime("%d.%m.%Y %H:%M:%S")


def _validate_gost_result(analytic: CalculationResult | None) -> CalculationResult:
    if analytic is None or analytic.calculation_basis != GOST_BASIS:
        raise ValueError(
            "Для отчёта «По ГОСТ» сначала выполните нормативный расчёт "
            "по ГОСТ 34758-2021 на странице 5."
        )
    return analytic


def _metric_map(analytic: CalculationResult) -> dict[str, object]:
    return {metric.key: metric for metric in analytic.metrics}


def _overall_status(analytic: CalculationResult) -> str:
    assessed = [
        metric
        for metric in analytic.metrics
        if metric.compliance is not ComplianceStatus.NOT_ASSESSED
    ]
    if assessed and all(
        metric.compliance is ComplianceStatus.COMPLIES for metric in assessed
    ):
        return ComplianceStatus.COMPLIES.value
    return ComplianceStatus.DOES_NOT_COMPLY.value


def _format_floor_marking(project: Project, floor_number: int) -> str:
    markers: list[str] = []
    for group in project.elevator_groups:
        markers.append(f"{group.name}: {'X' if floor_number in group.served_floors else 'I'}")
    return "; ".join(markers)


def build_gost_docx_report(
    project: Project,
    analytic: CalculationResult | None,
) -> bytes:
    """Формирует DOCX по составу сведений раздела 9 ГОСТ 34758-2021."""

    result = _validate_gost_result(analytic)
    metrics = _metric_map(result)
    group = project.group(result.group_id)
    elevator = group.elevators[0]
    served = [floor for floor in project.floors if floor.number in group.served_floors]
    travel_height = max(floor.elevation_m for floor in served) - min(
        floor.elevation_m for floor in served
    )

    document = Document()
    _configure_document(document)
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    header = section.header.paragraphs[0]
    header.text = "ГОСТ 34758-2021  |  ОТЧЁТ О ВЕРТИКАЛЬНОМ ТРАНСПОРТЕ"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("5B6770")

    title = document.add_paragraph(style="Title")
    title.add_run("Отчёт о результатах проектирования\nвертикального транспорта")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("Форма по ГОСТ 34758-2021, раздел 9")
    document.add_paragraph(project.metadata.name)
    document.add_paragraph(project.metadata.address or "Адрес не указан")

    status_paragraph = document.add_paragraph()
    status_run = status_paragraph.add_run(f"Результат оценки: {_overall_status(result)}")
    status_run.bold = True
    status_run.font.color.rgb = RGBColor.from_string(
        "1B5E20" if _overall_status(result) == ComplianceStatus.COMPLIES.value else "B3261E"
    )
    document.add_paragraph(f"Метод: {result.method}.")
    document.add_paragraph(
        f"Дата и время расчёта: "
        f"{_format_report_datetime(result.audit.calculated_at)}."
    )
    document.add_paragraph(
        "Состав отчёта соответствует перечню сведений раздела 9."
    )

    document.add_page_break()
    document.add_heading("1. Общие сведения", level=1)
    _add_table(
        document,
        ["Сведения по п. 9 а)", "Значение"],
        [
            [
                "Компания, выполнившая проектирование",
                project.metadata.designer or "не указана",
            ],
            ["Исполнитель проекта", project.metadata.calculation_author or "не указан"],
            ["Метод проектирования (анализа)", result.method],
            [
                "Дата и время оформления отчёта",
                _format_report_datetime(result.audit.calculated_at),
            ],
        ],
        [2.8, 4.1],
        show_header=False,
    )

    document.add_heading("2. Информация о здании", level=1)
    _add_table(
        document,
        ["Сведения по п. 9 б)", "Значение"],
        [
            ["Наименование здания", project.metadata.name],
            ["Адрес", project.metadata.address or "не указан"],
            ["Заказчик", project.metadata.customer or "не указан"],
            ["Назначение здания", project.building.building_type.value],
            ["Количество этажей", len(project.floors)],
            ["Количество лифтовых групп", len(project.elevator_groups)],
        ],
        [2.8, 4.1],
        show_header=False,
    )

    document.add_heading("2.1. Этажи и обслуживание лифтами", level=2)
    document.add_paragraph(
        f"Коэффициент заселённости: {project.building.occupancy_percent}%; "
        f"расчётное население здания: {project.population} чел."
    )
    _add_table(
        document,
        [
            "Этаж",
            "Маркировка",
            "Отметка, м",
            "Высота, м",
            "Назначение",
            "Расч. население",
        ],
        [
            [
                floor.number,
                floor.label or str(floor.number),
                f"{floor.elevation_m:.2f}",
                f"{floor.floor_height_m:.2f}",
                floor.purpose,
                f"{project.effective_floor_population(floor):.1f}",
            ]
            for floor in sorted(project.floors, key=lambda item: item.elevation_m, reverse=True)
        ],
        [0.65, 1.1, 1.0, 0.85, 1.75, 1.55],
    )

    document.add_heading("2.2. Размещение лифтовых групп", level=2)
    _add_table(
        document,
        ["Группа", "Расположение / зона", "Основной этаж", "Этажи"],
        [
            [
                item.name,
                item.service_zone_name,
                item.main_floor,
                ", ".join(map(str, item.served_floors)),
            ]
            for item in project.elevator_groups
        ],
        [1.1, 1.7, 1.0, 3.1],
    )

    document.add_heading("3. Критерии проектирования", level=1)
    _add_table(
        document,
        ["Параметр", "Принято", "Требование", "Статус"],
        [
            [
                "Провозная способность для пикового входящего потока",
                f"{metrics['handling_capacity_5min'].value:.2f} пасс./5 мин",
                metrics["handling_capacity_5min"].target_description or "",
                metrics["handling_capacity_5min"].compliance.value,
            ],
            [
                "Провозная способность к заселённости",
                f"{metrics['specific_capacity'].value:.2f} %/5 мин",
                metrics["specific_capacity"].target_description or "",
                metrics["specific_capacity"].compliance.value,
            ],
            [
                "Интервал движения",
                f"{metrics['interval'].value:.2f} с",
                metrics["interval"].target_description or "",
                metrics["interval"].compliance.value,
            ],
            [
                "Время движения на всю высоту",
                f"{metrics['full_height_time'].value:.2f} с",
                metrics["full_height_time"].target_description or "",
                metrics["full_height_time"].compliance.value,
            ],
        ],
        [2.8, 1.35, 1.55, 1.2],
    )

    document.add_heading("4. Информация о лифтовой установке", level=1)
    _add_table(
        document,
        ["Параметр", "Значение"],
        [
            ["Расположение группы", group.service_zone_name],
            ["Количество лифтов в группе", group.elevator_count],
            ["Высота подъёма", f"{travel_height:.2f} м"],
            [
                "Номинальная вместимость кабины",
                f"{metrics['nominal_capacity'].value:.0f} пасс.",
            ],
            ["Номинальная грузоподъёмность", f"{elevator.capacity_kg:.0f} кг"],
            [
                "Расчётная вместимость Pк",
                f"{metrics['actual_car_passengers'].value:.0f} пасс.",
            ],
            [
                "Правило округления Pк",
                "до ближайшего целого пассажира; 0,5 округляется вверх",
            ],
            ["Номинальная скорость", f"{elevator.speed_mps:.2f} м/с"],
            ["Ускорение", f"{elevator.acceleration_mps2:.2f} м/с²"],
            ["Замедление", f"{elevator.deceleration_mps2:.2f} м/с²"],
            ["Рывок", f"{elevator.jerk_mps3:.2f} м/с³"],
            [
                "Достижение номинальной скорости на межэтажном пролёте",
                (
                    "достигается"
                    if metrics["adjacent_floor_peak_speed"].value
                    >= elevator.speed_mps - 1e-9
                    else "не достигается"
                ),
            ],
            [
                "Максимальная скорость на межэтажном пролёте",
                f"{metrics['adjacent_floor_peak_speed'].value:.2f} м/с",
            ],
            [
                "Межэтажное время движения с разгоном и торможением",
                f"{metrics['adjacent_floor_profile_time'].value:.3f} с",
            ],
            [
                "Время входа или выхода пассажира",
                f"{next(trace.result for trace in result.formulas if trace.formula_id == 'gost_passenger_transfer_time'):.2f} с",
            ],
            ["Ширина дверного проёма", f"{elevator.door_width_m * 1000:.0f} мм"],
            ["Тип открывания дверей", elevator.door_opening_type.value],
            ["Время открывания дверей", f"{elevator.door_open_time_s:.2f} с"],
            ["Время закрывания дверей", f"{elevator.door_close_time_s:.2f} с"],
            ["Время предварительного открывания", f"{elevator.pre_open_time_s:.2f} с"],
            ["Время задержки закрывания", f"{elevator.door_dwell_time_s:.2f} с"],
            ["Время задержки начала движения", f"{elevator.start_brake_allowance_s:.2f} с"],
            ["Время, затрачиваемое на остановку", f"{metrics['stop_time'].value:.2f} с"],
        ],
        [3.4, 3.5],
    )

    if "parking_round_trip_addition" in metrics:
        document.add_heading(
            "4.1. Учёт паркинга — инженерное расширение (не формулы ГОСТ)",
            level=2,
        )
        document.add_paragraph(
            "Поправка учитывает альтернативный входящий поток с подземных "
            "этажей. Для консервативного расчёта принято, что каждый круговой "
            "рейс включает паркинг. Это не формула ГОСТ 34758-2021; фактическая "
            "доля потока учитывается в предварительном расчёте и симуляции."
        )
        _add_table(
            document,
            ["Параметр", "Значение"],
            [
                ["Доля входящего потока с паркинга", f"{metrics['parking_share'].value:.2f}%"],
                ["Допущение расчёта по ГОСТ", "паркинг учитывается в каждом рейсе"],
                ["Принятая вероятность заезда", f"{metrics['parking_trip_probability'].value:.2f}%"],
                ["Расчётный нижний уровень", f"{metrics['parking_lower_reversal'].value:.2f}"],
                ["Вероятное число остановок", f"{metrics['parking_probable_stops'].value:.2f}"],
                ["Ожидаемая глубина заезда", f"{metrics['parking_expected_depth'].value:.2f} м"],
                ["Круговой рейс по ГОСТ без паркинга", f"{metrics['gost_cycle_time_without_parking'].value:.2f} с"],
                ["Инженерная поправка паркинга", f"{metrics['parking_round_trip_addition'].value:.2f} с"],
                ["Круговой рейс с поправкой", f"{metrics['cycle_time'].value:.2f} с"],
            ],
            [3.4, 3.5],
        )

    document.add_heading("5. Расчётные данные провозной способности", level=1)
    _add_table(
        document,
        ["Параметр", "Результат", "Критерий", "Оценка"],
        [
            ["Этаж реверса", f"{metrics['highest_reversal'].value:.3f}", "—", "—"],
            ["Вероятное число остановок", f"{metrics['probable_stops'].value:.3f}", "—", "—"],
            ["Время кругового рейса", f"{metrics['cycle_time'].value:.3f} с", "—", "—"],
            [
                "Интервал в пиковый период",
                f"{metrics['interval'].value:.3f} с",
                metrics["interval"].target_description or "",
                metrics["interval"].compliance.value,
            ],
            [
                "Провозная способность за 5 мин",
                f"{metrics['handling_capacity_5min'].value:.3f} пасс.",
                metrics["handling_capacity_5min"].target_description or "",
                metrics["handling_capacity_5min"].compliance.value,
            ],
            [
                "Провозная способность к заселённости",
                f"{metrics['specific_capacity'].value:.3f} %",
                metrics["specific_capacity"].target_description or "",
                metrics["specific_capacity"].compliance.value,
            ],
            [
                "Заполнение кабины",
                f"{elevator.load_factor:.1%}",
                "обычно 0,8 по п. 6.5.3",
                "принято в расчёте",
            ],
            [
                "Итоговая оценка соответствия",
                _overall_status(result),
                "таблицы 1 и 4",
                _overall_status(result),
            ],
        ],
        [2.75, 1.45, 1.55, 1.15],
    )

    document.add_heading("6. Вывод и ограничения применимости", level=1)
    document.add_paragraph(
        f"По результатам расчётного метода выбранная конфигурация: {_overall_status(result).lower()} "
        "критериям проектирования ГОСТ 34758-2021 в пределах принятой расчётной модели."
    )
    for recommendation in result.recommendations:
        _add_list_item(
            document,
            f"{recommendation.problem} {recommendation.proposed_action} "
            f"Ограничение: {recommendation.limitations}",
        )
    for limitation in (
        "Расчёт выполнен для пикового пассажиропотока вверх при одном нижнем посадочном этаже.",
        "Этажи приняты равномерно заселёнными, лифтовая группа — однородной.",
        "Ориентировочное время ожидания не использовано как нормативный критерий.",
        "Перед договорным применением требуется контроль по экземпляру стандарта заказчика.",
    ):
        _add_list_item(document, limitation)

    document.add_page_break()
    document.add_heading("Приложение 1. Формулы и подстановка значений", level=1)
    for trace in result.formulas:
        add_word_formula_block(document, trace, heading_level=2)

    document.add_heading("Идентификация расчёта", level=1)
    _add_table(
        document,
        ["Параметр", "Значение"],
        [
            ["Версия приложения", __version__],
            ["Версия нормативной конфигурации", result.audit.configuration_version],
            ["Хэш исходных данных", result.audit.project_hash],
            ["Исполнитель", project.metadata.calculation_author or "не указан"],
            ["Подпись / дата", "____________________________"],
        ],
        [2.5, 4.4],
        add_spacing_after=False,
        font_size_pt=8.0,
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        f"ГОСТ 34758-2021 • {project.metadata.name} • Lift Traffic Analyzer {__version__}"
    )
    _set_document_language(document)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def build_gost_pdf_report(
    project: Project,
    analytic: CalculationResult | None,
) -> bytes:
    """Формирует компактную PDF-версию отчёта по разделу 9 ГОСТ."""

    result = _validate_gost_result(analytic)
    metrics = _metric_map(result)
    group = project.group(result.group_id)
    elevator = group.elevators[0]
    served = [floor for floor in project.floors if floor.number in group.served_floors]
    travel_height = max(floor.elevation_m for floor in served) - min(
        floor.elevation_m for floor in served
    )
    regular, bold = _register_fonts()
    stream = BytesIO()
    document = SimpleDocTemplate(
        stream,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=16 * mm,
        title=f"Отчёт по ГОСТ 34758-2021 — {project.metadata.name}",
        author=project.metadata.calculation_author,
    )
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "GOST Title",
            parent=base["Title"],
            fontName=bold,
            fontSize=19,
            leading=23,
            textColor=colors.HexColor("#1F4D78"),
            alignment=TA_LEFT,
            spaceAfter=10,
        ),
        "h1": ParagraphStyle(
            "GOST H1",
            parent=base["Heading1"],
            fontName=bold,
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=10,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "GOST Body",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=5,
        ),
        "small": ParagraphStyle(
            "GOST Small",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=7.5,
            leading=9,
        ),
        "formula_label": ParagraphStyle(
            "GOST Formula Label",
            parent=base["BodyText"],
            fontName=bold,
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#5B6770"),
            spaceAfter=2,
        ),
        "formula": ParagraphStyle(
            "GOST Formula",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=12.5,
            leading=18,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#111827"),
            spaceBefore=2,
            spaceAfter=7,
        ),
        "formula_legend": ParagraphStyle(
            "GOST Formula Legend",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#374151"),
            spaceBefore=0,
            spaceAfter=6,
        ),
        "formula_source": ParagraphStyle(
            "GOST Formula Source",
            parent=base["BodyText"],
            fontName=regular,
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#6B7280"),
            spaceAfter=10,
        ),
        "table_header": ParagraphStyle(
            "GOST Table Header",
            parent=base["BodyText"],
            fontName=bold,
            fontSize=7.5,
            leading=9,
            textColor=colors.white,
        ),
    }

    def paragraph(text: object, style: str = "body", markup: bool = False) -> Paragraph:
        content = str(text) if markup else escape(str(text))
        return Paragraph(content, styles[style])

    def table(
        headers: list[str],
        rows: list[list[object]],
        widths: list[float],
        *,
        show_header: bool = True,
    ) -> Table:
        def table_value(value: object) -> Paragraph:
            text = str(value).strip()
            color = {
                ComplianceStatus.COMPLIES.value: COMPLIES_COLOR,
                ComplianceStatus.DOES_NOT_COMPLY.value: DOES_NOT_COMPLY_COLOR,
            }.get(text)
            if color is None:
                return paragraph(value, "small")
            return paragraph(
                f'<font color="#{color}"><b>{escape(text)}</b></font>',
                "small",
                markup=True,
            )

        data: list[list[Paragraph]] = []
        if show_header:
            data.append([paragraph(item, "table_header") for item in headers])
        data.extend([[table_value(item) for item in row] for row in rows])
        output = Table(
            data,
            colWidths=[value * mm for value in widths],
            repeatRows=1 if show_header else 0,
            hAlign="LEFT",
        )
        commands = [
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#A9BBC4")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            (
                "ROWBACKGROUNDS",
                (0, 1 if show_header else 0),
                (-1, -1),
                [colors.white, colors.HexColor("#F4F7F9")],
            ),
        ]
        if show_header:
            commands.insert(
                0,
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
            )
        output.setStyle(TableStyle(commands))
        return output

    parking_story: list[object] = []
    if "parking_round_trip_addition" in metrics:
        parking_story = [
            paragraph(
                "4.1. Учёт паркинга — инженерное расширение (не формулы ГОСТ)",
                "h1",
            ),
            paragraph(
                "Поправка учитывает альтернативный входящий поток с подземных "
                "этажей. Для консервативного расчёта принято, что каждый круговой "
                "рейс включает паркинг. Это не формула ГОСТ 34758-2021; фактическая "
                "доля потока учитывается в предварительном расчёте и симуляции."
            ),
            table(
                ["Параметр", "Значение"],
                [
                    ["Доля входящего потока с паркинга", f"{metrics['parking_share'].value:.2f}%"],
                    ["Допущение расчёта по ГОСТ", "паркинг учитывается в каждом рейсе"],
                    ["Принятая вероятность заезда", f"{metrics['parking_trip_probability'].value:.2f}%"],
                    ["Расчётный нижний уровень", f"{metrics['parking_lower_reversal'].value:.2f}"],
                    ["Вероятное число остановок", f"{metrics['parking_probable_stops'].value:.2f}"],
                    ["Ожидаемая глубина заезда", f"{metrics['parking_expected_depth'].value:.2f} м"],
                    ["Круговой рейс по ГОСТ без паркинга", f"{metrics['gost_cycle_time_without_parking'].value:.2f} с"],
                    ["Инженерная поправка паркинга", f"{metrics['parking_round_trip_addition'].value:.2f} с"],
                    ["Круговой рейс с поправкой", f"{metrics['cycle_time'].value:.2f} с"],
                ],
                [90, 78],
            ),
        ]

    story: list[object] = [
        paragraph(
            "Отчёт о результатах проектирования<br/>вертикального транспорта",
            "title",
            markup=True,
        ),
        paragraph("Форма по ГОСТ 34758-2021, раздел 9"),
        Spacer(1, 5 * mm),
        table(
            ["Поле", "Значение"],
            [
                ["Объект", project.metadata.name],
                ["Адрес", project.metadata.address or "не указан"],
                ["Заказчик", project.metadata.customer or "не указан"],
                ["Исполнитель", project.metadata.calculation_author or "не указан"],
                ["Метод", result.method],
                ["Итог", _overall_status(result)],
            ],
            [52, 116],
        ),
        PageBreak(),
        paragraph("1. Общие сведения", "h1"),
        table(
            ["Параметр", "Значение"],
            [
                ["Компания", project.metadata.designer or "не указана"],
                ["Исполнитель", project.metadata.calculation_author or "не указан"],
                ["Метод анализа", result.method],
                [
                    "Дата и время",
                    _format_report_datetime(result.audit.calculated_at),
                ],
            ],
            [56, 112],
            show_header=False,
        ),
        paragraph("2. Информация о здании", "h1"),
        table(
            ["Параметр", "Значение"],
            [
                ["Наименование / адрес", f"{project.metadata.name}; {project.metadata.address or 'не указан'}"],
                ["Заказчик", project.metadata.customer or "не указан"],
                ["Назначение", project.building.building_type.value],
                ["Этажей / групп", f"{len(project.floors)} / {len(project.elevator_groups)}"],
                ["Коэффициент заселённости", f"{project.building.occupancy_percent}%"],
                ["Расчётное население", f"{project.population} чел."],
            ],
            [56, 112],
            show_header=False,
        ),
        table(
            ["Этаж", "Метка", "Высота", "Назначение", "Расч. население"],
            [
                [
                    floor.number,
                    floor.label or floor.number,
                    f"{floor.floor_height_m:.2f} м",
                    floor.purpose,
                    f"{project.effective_floor_population(floor):.1f}",
                ]
                for floor in sorted(project.floors, key=lambda item: item.elevation_m, reverse=True)
            ],
            [17, 23, 28, 60, 40],
        ),
        paragraph("3. Критерии проектирования", "h1"),
        table(
            ["Параметр", "Факт", "Требование", "Оценка"],
            [
                [
                    "Провозная способность",
                    f"{metrics['handling_capacity_5min'].value:.2f} пасс./5 мин",
                    metrics["handling_capacity_5min"].target_description or "",
                    metrics["handling_capacity_5min"].compliance.value,
                ],
                [
                    "Интервал",
                    f"{metrics['interval'].value:.2f} с",
                    metrics["interval"].target_description or "",
                    metrics["interval"].compliance.value,
                ],
                [
                    "Время на всю высоту",
                    f"{metrics['full_height_time'].value:.2f} с",
                    metrics["full_height_time"].target_description or "",
                    metrics["full_height_time"].compliance.value,
                ],
            ],
            [55, 38, 40, 35],
        ),
        paragraph("4. Информация о лифтовой установке", "h1"),
        table(
            ["Параметр", "Значение"],
            [
                ["Группа / зона", f"{group.name}; {group.service_zone_name}"],
                ["Количество лифтов", group.elevator_count],
                ["Высота подъёма", f"{travel_height:.2f} м"],
                ["Грузоподъёмность / вместимость", f"{elevator.capacity_kg:.0f} кг / {metrics['nominal_capacity'].value:.0f} пасс."],
                ["Расчётная вместимость Pк", f"{metrics['actual_car_passengers'].value:.0f} пасс."],
                ["Правило округления Pк", "до ближайшего целого; 0,5 вверх"],
                ["Скорость / дверь", f"{elevator.speed_mps:.2f} м/с / {elevator.door_width_m * 1000:.0f} мм"],
                [
                    "Ускорение / замедление / рывок",
                    (
                        f"{elevator.acceleration_mps2:.2f} м/с² / "
                        f"{elevator.deceleration_mps2:.2f} м/с² / "
                        f"{elevator.jerk_mps3:.2f} м/с³"
                    ),
                ],
                [
                    "Номинальная скорость на межэтажном пролёте",
                    (
                        "достигается"
                        if metrics["adjacent_floor_peak_speed"].value
                        >= elevator.speed_mps - 1e-9
                        else "не достигается"
                    ),
                ],
                [
                    "Максимальная скорость на межэтажном пролёте",
                    f"{metrics['adjacent_floor_peak_speed'].value:.2f} м/с",
                ],
                [
                    "Межэтажное время с разгоном и торможением",
                    f"{metrics['adjacent_floor_profile_time'].value:.3f} с",
                ],
                ["Открытие / закрытие / предв. открытие", f"{elevator.door_open_time_s:.2f} / {elevator.door_close_time_s:.2f} / {elevator.pre_open_time_s:.2f} с"],
                ["Тип открывания дверей", elevator.door_opening_type.value],
                ["Задержка закрытия / пуска", f"{elevator.door_dwell_time_s:.2f} / {elevator.start_brake_allowance_s:.2f} с"],
                ["Время остановки", f"{metrics['stop_time'].value:.2f} с"],
            ],
            [76, 92],
        ),
        *parking_story,
        paragraph("5. Расчётные данные провозной способности", "h1"),
        table(
            ["Параметр", "Результат", "Оценка"],
            [
                ["Этаж реверса", f"{metrics['highest_reversal'].value:.3f}", "—"],
                ["Вероятное число остановок", f"{metrics['probable_stops'].value:.3f}", "—"],
                ["Время кругового рейса", f"{metrics['cycle_time'].value:.3f} с", "—"],
                ["Интервал", f"{metrics['interval'].value:.3f} с", metrics["interval"].compliance.value],
                ["Провозная способность за 5 мин", f"{metrics['handling_capacity_5min'].value:.3f}", metrics["handling_capacity_5min"].compliance.value],
                ["Провозная способность, %", f"{metrics['specific_capacity'].value:.3f}", metrics["specific_capacity"].compliance.value],
                ["Заполнение кабины", f"{elevator.load_factor:.1%}", "принято"],
                ["Итог", _overall_status(result), _overall_status(result)],
            ],
            [82, 48, 38],
        ),
        paragraph("6. Вывод и ограничения", "h1"),
        paragraph(
            f"Выбранная конфигурация: {_overall_status(result).lower()} критериям "
            "расчётного метода ГОСТ 34758-2021."
        ),
        paragraph(
            "Область применимости: восходящий пик, один нижний посадочный этаж, "
            "равномерное заселение и однородная группа. Перед договорным применением "
            "требуется контроль по экземпляру стандарта заказчика."
        ),
        paragraph("Формулы и ссылки", "h1"),
    ]
    for trace in result.formulas:
        source = trace.standard
        if trace.clause:
            source = f"{source}, {trace.clause}"
        story.append(
            KeepTogether(
                [
                    paragraph(f"<b>{escape(trace.title_ru)}</b>", markup=True),
                    paragraph("Расчётная зависимость", "formula_label"),
                    paragraph(pdf_formula_markup(trace), "formula", markup=True),
                    paragraph(
                        pdf_formula_legend_markup(trace),
                        "formula_legend",
                        markup=True,
                    ),
                    paragraph("Подстановка значений", "formula_label"),
                    paragraph(
                        pdf_substitution_markup(trace),
                        "formula",
                        markup=True,
                    ),
                    paragraph(
                        f"<b>Результат:</b> {trace.result:.2f} "
                        f"{escape(trace.unit)}<br/>Источник: {escape(source)}.",
                        "formula_source",
                        markup=True,
                    ),
                ]
            )
        )
    story.append(
        paragraph(
            f"Аудит: версия {__version__}; конфигурация {result.audit.configuration_version}; "
            f"хэш {result.audit.project_hash}."
        )
    )

    def footer(canvas: object, doc: object) -> None:
        canvas.saveState()
        canvas.setFont(regular, 8)
        canvas.setFillColor(colors.HexColor("#5B6770"))
        canvas.drawString(18 * mm, 9 * mm, "ГОСТ 34758-2021")
        canvas.drawRightString(A4[0] - 18 * mm, 9 * mm, f"Стр. {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return stream.getvalue()
