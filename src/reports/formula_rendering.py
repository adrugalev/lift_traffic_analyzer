"""Оформление расчётных формул в DOCX- и PDF-отчётах."""

from __future__ import annotations

from collections.abc import Sequence
from functools import lru_cache
import re
from typing import Any, TypeAlias
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from src.models.results import FormulaTrace
from src.reference_guide import FORMULA_VARIABLES_IN_ORDER
from src.services.configuration_service import ConfigurationService


MathSpec: TypeAlias = tuple[Any, ...]
WORD_LANGUAGE = "ru-RU"
_FORMULA_CONFIGURATION = ConfigurationService()


def _run(text: str, *, italic: bool = False) -> MathSpec:
    return ("run", text, italic)


def _seq(*parts: MathSpec) -> MathSpec:
    return ("seq", *parts)


def _sub(base: MathSpec, index: MathSpec) -> MathSpec:
    return ("sub", base, index)


def _sup(base: MathSpec, exponent: MathSpec) -> MathSpec:
    return ("sup", base, exponent)


def _frac(numerator: MathSpec, denominator: MathSpec) -> MathSpec:
    return ("frac", numerator, denominator)


def _delim(
    content: MathSpec,
    opening: str = "(",
    closing: str = ")",
) -> MathSpec:
    return ("delim", opening, closing, content)


def _rad(content: MathSpec) -> MathSpec:
    return ("rad", content)


def _sum(lower: MathSpec, upper: MathSpec, content: MathSpec) -> MathSpec:
    return ("sum", lower, upper, content)


def _v(name: str, index: str | None = None) -> MathSpec:
    variable = _run(name, italic=True)
    if index is None:
        return variable
    return _sub(variable, _run(index, italic=False))


def _n(value: object) -> MathSpec:
    return _run(str(value), italic=False)


def _op(value: str) -> MathSpec:
    return _run(value, italic=False)


def _formula_specs() -> dict[str, MathSpec]:
    one = _n("1")
    two = _n("2")
    three_hundred = _n("300")

    preview_stops = _seq(
        _v("N"),
        _delim(
            _seq(
                one,
                _op(" − "),
                _sup(
                    _delim(
                        _frac(
                            _seq(_v("N"), _op(" − "), one),
                            _v("N"),
                        )
                    ),
                    _v("C"),
                ),
            ),
            "[",
            "]",
        ),
    )
    preview_reversal = _sum(
        _seq(_v("k"), _op(" = "), one),
        _v("N"),
        _delim(
            _seq(
                one,
                _op(" − "),
                _sup(
                    _delim(
                        _frac(
                            _seq(_v("k"), _op(" − "), one),
                            _v("N"),
                        )
                    ),
                    _v("C"),
                ),
            ),
            "[",
            "]",
        ),
    )
    gost_stops = _seq(
        _v("N", "эт"),
        _delim(
            _seq(
                one,
                _op(" − "),
                _sup(
                    _delim(
                        _seq(
                            one,
                            _op(" − "),
                            _frac(one, _v("N", "эт")),
                        )
                    ),
                    _v("P", "к"),
                ),
            ),
            "[",
            "]",
        ),
    )
    gost_reversal_sum = _sum(
        _seq(_v("i"), _op(" = "), one),
        _seq(_v("N", "эт"), _op(" − "), one),
        _sup(
            _delim(_frac(_v("i"), _v("N", "эт"))),
            _v("P", "к"),
        ),
    )

    return {
        "calculated_car_capacity": _seq(
            _v("C"),
            _op(" = min"),
            _delim(
                _seq(
                    _v("C", "ном"),
                    _op(", "),
                    _sub(_op("окр"), _op("0,5↑")),
                    _delim(_seq(_v("C", "ном"), _op(" · "), _v("k", "зап"))),
                )
            ),
        ),
        "probable_stops_uniform": _seq(
            _v("S"),
            _op(" = "),
            preview_stops,
        ),
        "highest_reversal_uniform": _seq(
            _v("H"),
            _op(" = "),
            preview_reversal,
        ),
        "motion_time": _seq(
            _v("t", "дв"),
            _op(" = "),
            _v("f"),
            _delim(
                _seq(
                    _v("d"),
                    _op(", "),
                    _v("v"),
                    _op(", "),
                    _v("a"),
                    _op(", "),
                    _v("b"),
                )
            ),
            _op(" = "),
            _v("t", "разг"),
            _op(" + "),
            _v("t", "уст"),
            _op(" + "),
            _v("t", "торм"),
        ),
        "engineering_cycle_time": _seq(
            _v("T", "цикл"),
            _op(" = "),
            two,
            _op(" · "),
            _v("t", "дв"),
            _op(" + "),
            _v("S"),
            _op(" · "),
            _v("t", "ост"),
            _op(" + "),
            _v("C"),
            _delim(
                _seq(
                    _v("t", "пос"),
                    _op(" + "),
                    _v("t", "выс"),
                )
            ),
        ),
        "interval": _seq(
            _v("I"),
            _op(" = "),
            _frac(_v("T", "цикл"), _v("L")),
        ),
        "handling_capacity_5min": _seq(
            _v("HC", "5"),
            _op(" = "),
            _frac(
                _seq(three_hundred, _op(" · "), _v("C")),
                _v("I"),
            ),
        ),
        "average_wait_proxy": _seq(
            _v("t", "ож, ор"),
            _op(" ≈ "),
            _frac(_v("I"), two),
        ),
        "gost_nominal_capacity": _seq(
            _v("P", "ном"),
            _op(" = "),
            _sub(_op("окр"), _op("0,5↑")),
            _delim(_frac(_v("Q"), _n("75"))),
        ),
        "gost_calculated_car_capacity": _seq(
            _v("P", "к"),
            _op(" = "),
            _sub(_op("окр"), _op("0,5↑")),
            _delim(_seq(_v("P", "ном"), _op(" · "), _v("k", "з"))),
        ),
        "gost_passenger_transfer_time": _seq(
            _v("t", "в"),
            _op(" = табл. 3"),
            _delim(_v("b", "двери")),
        ),
        "gost_probable_stops": _seq(
            _v("S"),
            _op(" = "),
            gost_stops,
        ),
        "gost_reversal_floor": _seq(
            _v("N", "р"),
            _op(" = "),
            _v("N", "эт"),
            _op(" − "),
            gost_reversal_sum,
        ),
        "gost_adjacent_floor_time": _seq(
            _v("t", "эт.н"),
            _op(" = "),
            _frac(_v("h", "эт"), _v("v", "н")),
        ),
        "gost_adjacent_floor_profile_time": _seq(
            _v("t", "эт"),
            _op(" = "),
            _sub(_v("f"), _op("S")),
            _delim(
                _seq(
                    _v("h", "эт"),
                    _op(", "),
                    _v("v", "н"),
                    _op(", "),
                    _v("a"),
                    _op(", "),
                    _v("b"),
                    _op(", "),
                    _v("j"),
                )
            ),
            _op(" = "),
            _v("t", "разг"),
            _op(" + "),
            _v("t", "уст"),
            _op(" + "),
            _v("t", "торм"),
        ),
        "gost_stop_time": _seq(
            _v("t", "ост"),
            _op(" = "),
            _v("t", "з"),
            _op(" + "),
            _v("t", "з.д"),
            _op(" + "),
            _v("t", "эт"),
            _op(" − "),
            _v("t", "пр"),
            _op(" + "),
            _v("t", "о"),
            _op(" + "),
            _v("t", "з.з"),
            _op(" − "),
            _v("t", "эт.н"),
        ),
        "gost_round_trip_time": _seq(
            _v("T"),
            _op(" = "),
            two,
            _op(" · "),
            _v("N", "р"),
            _op(" · "),
            _v("t", "эт.н"),
            _op(" + "),
            _delim(_seq(_v("S"), _op(" + "), one)),
            _op(" · "),
            _v("t", "ост"),
            _op(" + "),
            two,
            _op(" · "),
            _v("P", "к"),
            _op(" · "),
            _v("t", "в"),
        ),
        "gost_interval": _seq(
            _v("t", "и"),
            _op(" = "),
            _frac(_v("T"), _v("N", "л")),
        ),
        "gost_group_handling_capacity": _seq(
            _v("P", "5"),
            _op(" = "),
            _frac(
                _seq(
                    three_hundred,
                    _op(" · "),
                    _v("P", "к"),
                    _op(" · "),
                    _v("N", "л"),
                ),
                _v("T"),
            ),
            _op(" = "),
            _frac(
                _seq(three_hundred, _op(" · "), _v("P", "к")),
                _v("t", "и"),
            ),
        ),
        "gost_handling_capacity_percent": _seq(
            _op("%"),
            _v("P", "5"),
            _op(" = "),
            _frac(
                _seq(_v("P", "5"), _op(" · "), _n("100")),
                _v("A"),
            ),
        ),
        "parking_lower_reversal": _seq(
            _v("H", "м"),
            _op(" = "),
            _sum(
                _seq(_v("k"), _op(" = "), one),
                _v("M"),
                _delim(
                    _seq(
                        one,
                        _op(" − "),
                        _sup(
                            _delim(
                                _seq(
                                    one,
                                    _op(" − "),
                                    _v("q", "м"),
                                    _op(" · "),
                                    _frac(
                                        _seq(_v("M"), _op(" − "), _v("k"), _op(" + "), one),
                                        _v("M"),
                                    ),
                                )
                            ),
                            _v("P", "к"),
                        ),
                    ),
                    "[",
                    "]",
                ),
            ),
        ),
        "parking_expected_depth": _seq(
            _v("D", "м"),
            _op(" = "),
            _sum(
                _seq(_v("k"), _op(" = "), one),
                _v("M"),
                _seq(
                    _v("Δh", "к"),
                    _op(" · "),
                    _delim(
                        _seq(
                            one,
                            _op(" − "),
                            _sup(
                                _delim(
                                    _seq(
                                        one,
                                        _op(" − "),
                                        _v("q", "м"),
                                        _op(" · "),
                                        _frac(
                                            _seq(_v("M"), _op(" − "), _v("k"), _op(" + "), one),
                                            _v("M"),
                                        ),
                                    )
                                ),
                                _v("P", "к"),
                            ),
                        ),
                        "[",
                        "]",
                    ),
                ),
            ),
        ),
        "parking_probable_stops": _seq(
            _v("S", "м"),
            _op(" = "),
            _v("M"),
            _op(" · "),
            _delim(
                _seq(
                    one,
                    _op(" − "),
                    _sup(
                        _delim(
                            _seq(
                                one,
                                _op(" − "),
                                _frac(_v("q", "м"), _v("M")),
                            )
                        ),
                        _v("P", "к"),
                    ),
                ),
                "[",
                "]",
            ),
        ),
        "parking_round_trip_extension": _seq(
            _v("T", "м"),
            _op(" = "),
            _v("T", "ГОСТ"),
            _op(" + "),
            _frac(_seq(two, _op(" · "), _v("D", "м")), _v("v", "н")),
            _op(" + "),
            _v("S", "м"),
            _op(" · "),
            _v("t", "ост.м"),
        ),
        "gost_nominal_speed": _seq(
            _v("v", "н"),
            _op(" = "),
            _frac(_v("H", "max"), _v("t", "н")),
        ),
        # Варианты формул из отчёта-образца. Они готовы к использованию,
        # если кинематическая трассировка будет добавлена в расчётный движок.
        "kinematic_acceleration_distance": _seq(
            _v("s", "vm"),
            _op(" = "),
            _frac(_sup(_v("v", "m"), two), _seq(two, _v("a"))),
            _op(" + "),
            _frac(_seq(_v("a"), _op(" · "), _v("v", "m")), _seq(two, _v("j"))),
            _op(";  "),
            _seq(two, _v("s", "vm")),
            _op(" ≤ "),
            _v("d", "f"),
        ),
        "kinematic_maximum_speed": _seq(
            _v("v", "m"),
            _op(" = −"),
            _frac(_sup(_v("a"), two), _seq(two, _v("j"))),
            _op(" + "),
            _rad(
                _seq(
                    _v("a"),
                    _op(" · "),
                    _v("d", "f"),
                    _op(" + "),
                    _sup(
                        _delim(
                            _frac(_sup(_v("a"), two), _seq(two, _v("j")))
                        ),
                        two,
                    ),
                )
            ),
        ),
        "kinematic_adjacent_floor_time": _seq(
            _v("t", "f"),
            _delim(one),
            _op(" = "),
            _frac(_v("d", "f"), _v("v", "m")),
            _op(" + "),
            _frac(_v("a"), _v("j")),
            _op(" + "),
            _frac(_v("v", "m"), _v("a")),
        ),
    }


FORMULA_SPECS = _formula_specs()


PDF_FORMULAS = {
    "calculated_car_capacity": (
        "C = min(C<sub>ном</sub>, окр<sub>0,5↑</sub>"
        "(C<sub>ном</sub> · k<sub>зап</sub>))"
    ),
    "probable_stops_uniform": (
        "S = N [1 − ((N − 1) ⁄ N)<super>C</super>]"
    ),
    "highest_reversal_uniform": (
        "H = Σ<sub>k=1</sub><super>N</super> "
        "[1 − ((k − 1) ⁄ N)<super>C</super>]"
    ),
    "motion_time": (
        "t<sub>дв</sub> = f(d, v, a, b) = "
        "t<sub>разг</sub> + t<sub>уст</sub> + t<sub>торм</sub>"
    ),
    "engineering_cycle_time": (
        "T<sub>цикл</sub> = 2 · t<sub>дв</sub> + S · t<sub>ост</sub> "
        "+ C (t<sub>пос</sub> + t<sub>выс</sub>)"
    ),
    "interval": "I = T<sub>цикл</sub> ⁄ L",
    "handling_capacity_5min": "HC<sub>5</sub> = 300 · C ⁄ I",
    "average_wait_proxy": "t<sub>ож, ор</sub> ≈ I ⁄ 2",
    "gost_nominal_capacity": (
        "P<sub>ном</sub> = окр<sub>0,5↑</sub>(Q ⁄ 75)"
    ),
    "gost_calculated_car_capacity": (
        "P<sub>к</sub> = окр<sub>0,5↑</sub>"
        "(P<sub>ном</sub> · k<sub>з</sub>)"
    ),
    "gost_passenger_transfer_time": (
        "t<sub>в</sub> = табл. 3(b<sub>двери</sub>)"
    ),
    "gost_probable_stops": (
        "S = N<sub>эт</sub> [1 − (1 − 1 ⁄ N<sub>эт</sub>)"
        "<super>P<sub>к</sub></super>]"
    ),
    "gost_reversal_floor": (
        "N<sub>р</sub> = N<sub>эт</sub> − "
        "Σ<sub>i=1</sub><super>N<sub>эт</sub>−1</super> "
        "(i ⁄ N<sub>эт</sub>)<super>P<sub>к</sub></super>"
    ),
    "gost_adjacent_floor_time": (
        "t<sub>эт.н</sub> = h<sub>эт</sub> ⁄ v<sub>н</sub>"
    ),
    "gost_adjacent_floor_profile_time": (
        "t<sub>эт</sub> = f<sub>S</sub>(h<sub>эт</sub>, v<sub>н</sub>, "
        "a, b, j) = t<sub>разг</sub> + t<sub>уст</sub> + t<sub>торм</sub>"
    ),
    "gost_stop_time": (
        "t<sub>ост</sub> = t<sub>з</sub> + t<sub>з.д</sub> "
        "+ t<sub>эт</sub> − t<sub>пр</sub> + t<sub>о</sub> "
        "+ t<sub>з.з</sub> − t<sub>эт.н</sub>"
    ),
    "gost_round_trip_time": (
        "T = 2 · N<sub>р</sub> · t<sub>эт.н</sub> "
        "+ (S + 1) · t<sub>ост</sub> + 2 · P<sub>к</sub> · t<sub>в</sub>"
    ),
    "gost_interval": "t<sub>и</sub> = T ⁄ N<sub>л</sub>",
    "gost_group_handling_capacity": (
        "P<sub>5</sub> = 300 · P<sub>к</sub> · N<sub>л</sub> ⁄ T "
        "= 300 · P<sub>к</sub> ⁄ t<sub>и</sub>"
    ),
    "gost_handling_capacity_percent": (
        "%P<sub>5</sub> = P<sub>5</sub> · 100 ⁄ A"
    ),
    "parking_lower_reversal": (
        "H<sub>м</sub> = Σ<sub>k=1</sub><super>M</super> "
        "[1 − (1 − q<sub>м</sub> · (M − k + 1) ⁄ M)"
        "<super>P<sub>к</sub></super>]"
    ),
    "parking_expected_depth": (
        "D<sub>м</sub> = Σ<sub>k=1</sub><super>M</super> Δh<sub>к</sub> · "
        "[1 − (1 − q<sub>м</sub> · (M − k + 1) ⁄ M)"
        "<super>P<sub>к</sub></super>]"
    ),
    "parking_probable_stops": (
        "S<sub>м</sub> = M · [1 − (1 − q<sub>м</sub> ⁄ M)"
        "<super>P<sub>к</sub></super>]"
    ),
    "parking_round_trip_extension": (
        "T<sub>м</sub> = T<sub>ГОСТ</sub> + 2 · D<sub>м</sub> ⁄ "
        "v<sub>н</sub> + S<sub>м</sub> · t<sub>ост.м</sub>"
    ),
    "gost_nominal_speed": (
        "v<sub>н</sub> = H<sub>max</sub> ⁄ t<sub>н</sub>"
    ),
    "kinematic_acceleration_distance": (
        "s<sub>вм</sub> = v<sub>м</sub><super>2</super> ⁄ (2a) "
        "+ a · v<sub>м</sub> ⁄ (2j); "
        "2s<sub>вм</sub> ≤ d<sub>ф</sub>"
    ),
    "kinematic_maximum_speed": (
        "v<sub>м</sub> = −a<super>2</super> ⁄ (2j) + √("
        "a · d<sub>ф</sub> + (a<super>2</super> ⁄ (2j))<super>2</super>)"
    ),
}


def _append_math(parent: Any, spec: MathSpec) -> None:
    kind = spec[0]
    if kind == "seq":
        for part in spec[1:]:
            _append_math(parent, part)
        return
    if kind == "run":
        element = OxmlElement("m:r")
        math_properties = OxmlElement("m:rPr")
        style = OxmlElement("m:sty")
        style.set(qn("m:val"), "i" if spec[2] else "p")
        math_properties.append(style)
        element.append(math_properties)
        word_properties = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), "Cambria Math")
        word_properties.append(fonts)
        language = OxmlElement("w:lang")
        for attribute in ("val", "eastAsia", "bidi"):
            language.set(qn(f"w:{attribute}"), WORD_LANGUAGE)
        word_properties.append(language)
        element.append(word_properties)
        text = OxmlElement("m:t")
        text.text = spec[1]
        element.append(text)
        parent.append(element)
        return
    if kind in {"sub", "sup"}:
        element = OxmlElement("m:sSub" if kind == "sub" else "m:sSup")
        base = OxmlElement("m:e")
        index = OxmlElement("m:sub" if kind == "sub" else "m:sup")
        _append_math(base, spec[1])
        _append_math(index, spec[2])
        element.extend((base, index))
        parent.append(element)
        return
    if kind == "frac":
        element = OxmlElement("m:f")
        numerator = OxmlElement("m:num")
        denominator = OxmlElement("m:den")
        _append_math(numerator, spec[1])
        _append_math(denominator, spec[2])
        element.extend((numerator, denominator))
        parent.append(element)
        return
    if kind == "delim":
        element = OxmlElement("m:d")
        properties = OxmlElement("m:dPr")
        opening = OxmlElement("m:begChr")
        opening.set(qn("m:val"), spec[1])
        closing = OxmlElement("m:endChr")
        closing.set(qn("m:val"), spec[2])
        properties.extend((opening, closing))
        content = OxmlElement("m:e")
        _append_math(content, spec[3])
        element.extend((properties, content))
        parent.append(element)
        return
    if kind == "rad":
        element = OxmlElement("m:rad")
        properties = OxmlElement("m:radPr")
        degree_hidden = OxmlElement("m:degHide")
        degree_hidden.set(qn("m:val"), "1")
        properties.append(degree_hidden)
        degree = OxmlElement("m:deg")
        content = OxmlElement("m:e")
        _append_math(content, spec[1])
        element.extend((properties, degree, content))
        parent.append(element)
        return
    if kind == "sum":
        element = OxmlElement("m:nary")
        properties = OxmlElement("m:naryPr")
        character = OxmlElement("m:chr")
        character.set(qn("m:val"), "∑")
        limit_location = OxmlElement("m:limLoc")
        limit_location.set(qn("m:val"), "undOvr")
        grow = OxmlElement("m:grow")
        grow.set(qn("m:val"), "1")
        properties.extend((character, limit_location, grow))
        lower = OxmlElement("m:sub")
        upper = OxmlElement("m:sup")
        content = OxmlElement("m:e")
        _append_math(lower, spec[1])
        _append_math(upper, spec[2])
        _append_math(content, spec[3])
        element.extend((properties, lower, upper, content))
        parent.append(element)
        return
    raise ValueError(f"Неизвестный тип математического узла: {kind}")


def _pretty_linear_text(text: str) -> str:
    replacements = (
        ("T_cycle", "Tцикл"),
        ("t_motion", "tдв"),
        ("t_cruise", "tравн"),
        ("t_board", "tвх"),
        ("t_alight", "tвых"),
        ("t_stop", "tост"),
        ("C_nom", "Cном"),
        ("k_fill", "kзап"),
        ("AWT_proxy", "tож, ор"),
        ("HC5", "HC₅"),
        ("Pном", "Pном"),
        ("Nэт", "Nэт"),
        ("Nр", "Nр"),
        ("Pк", "Pк"),
        ("Nл", "Nл"),
        ("tэт.н", "tэт.н"),
        ("tост", "tост"),
        ("tз.д", "tз.д"),
        ("tз.з", "tз.з"),
        ("tпр", "tпр"),
        ("tи", "tи"),
        ("Vн,min", "vн,min"),
        ("Vн", "vн"),
        ("Hmax", "Hmax"),
        (" / ", " ÷ "),
        (" * ", " × "),
        ("^", " ^ "),
        ("..", "…"),
    )
    result = text
    for old, new in replacements:
        result = result.replace(old, new)
    return " ".join(result.split())


def _formula_symbol_parts(symbol: str) -> tuple[str, str, str]:
    prefix = ""
    body = symbol
    if body.startswith("%"):
        prefix = "%"
        body = body[1:]
    base_length = 2 if body.startswith("HC") else 1
    return prefix, body[:base_length], body[base_length:]


def _formula_symbol_spec(symbol: str) -> MathSpec:
    prefix, base, index = _formula_symbol_parts(symbol)
    variable = _v(base, index or None)
    return _seq(_op(prefix), variable) if prefix else variable


@lru_cache(maxsize=1)
def _symbol_aliases() -> dict[str, str]:
    symbols = {
        symbol
        for values in FORMULA_VARIABLES_IN_ORDER.values()
        for symbol in values
    }
    aliases = {symbol: symbol for symbol in symbols}
    aliases.update(
        {
            "HC₅": "HC5",
            "tож, ор": "tож,ор",
            "fS": "fS",
        }
    )
    return aliases


def _format_substitution_numbers(text: str) -> str:
    """Приводит расчётные числа к двум знакам, не меняя номера таблиц и пределы сумм."""

    table_markers: dict[str, str] = {}

    def protect_table_number(match: re.Match[str]) -> str:
        marker = f"§TABLE_{len(table_markers)}§"
        table_markers[marker] = match.group(0)
        return marker

    protected = re.sub(r"табл\.\s*\d+", protect_table_number, text)
    number_pattern = re.compile(
        r"(?<![\w.])(-?\d+(?:[.,]\d+)?)(?![\w.]|\.\.)",
        flags=re.UNICODE,
    )

    def format_number(match: re.Match[str]) -> str:
        value = float(match.group(1).replace(",", "."))
        return f"{value:.2f}"

    formatted = number_pattern.sub(format_number, protected)
    for marker, original in table_markers.items():
        formatted = formatted.replace(marker, original)
    return formatted


def formatted_substitution_text(trace: FormulaTrace) -> str:
    """Возвращает подстановку с пользовательскими обозначениями и двумя знаками."""

    return _pretty_linear_text(
        _format_substitution_numbers(trace.substituted_expression)
    )


def _linear_math_tokens(text: str) -> list[tuple[str, str]]:
    aliases = _symbol_aliases()
    names = sorted(
        ("окр₀,₅↑", *aliases),
        key=len,
        reverse=True,
    )
    alternatives = "|".join(re.escape(name) for name in names)
    pattern = re.compile(
        rf"(?<![A-Za-zА-Яа-яЁё])(?:{alternatives})(?![A-Za-zА-Яа-яЁё])"
    )
    tokens: list[tuple[str, str]] = []
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            tokens.append(("text", text[cursor : match.start()]))
        value = match.group(0)
        if value == "окр₀,₅↑":
            tokens.append(("round", value))
        else:
            tokens.append(("symbol", aliases[value]))
        cursor = match.end()
    if cursor < len(text):
        tokens.append(("text", text[cursor:]))
    return tokens


def _linear_math_spec(text: str) -> MathSpec:
    parts: list[MathSpec] = []
    for kind, value in _linear_math_tokens(text):
        if kind == "symbol":
            parts.append(_formula_symbol_spec(value))
        elif kind == "round":
            parts.append(_sub(_op("окр"), _op("0,5↑")))
        else:
            parts.append(_run(value, italic=False))
    return _seq(*parts) if parts else _run("", italic=False)


@lru_cache(maxsize=None)
def formula_variable_rows(formula_id: str) -> tuple[tuple[str, str, str], ...]:
    """Возвращает полный упорядоченный перечень обозначений одной формулы."""

    definition = (
        _FORMULA_CONFIGURATION.formulas()
        .get("formulas", {})
        .get(formula_id, {})
    )
    variables = definition.get("variables", {})
    ordered = list(FORMULA_VARIABLES_IN_ORDER.get(formula_id, ()))
    ordered.extend(symbol for symbol in variables if symbol not in ordered)
    return tuple(
        (
            symbol,
            str(variables[symbol].get("title_ru", "")),
            str(variables[symbol].get("unit", "—")),
        )
        for symbol in ordered
        if symbol in variables
    )


def _set_word_run_font(run: Any, *, math: bool = False) -> None:
    font_name = "Cambria Math" if math else "Arial"
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    language = run._element.get_or_add_rPr().find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        run._element.get_or_add_rPr().append(language)
    for attribute in ("val", "eastAsia", "bidi"):
        language.set(qn(f"w:{attribute}"), WORD_LANGUAGE)


def _append_word_formula_symbol(paragraph: Any, symbol: str) -> None:
    prefix, base, index = _formula_symbol_parts(symbol)
    if prefix:
        prefix_run = paragraph.add_run(prefix)
        _set_word_run_font(prefix_run, math=True)
    base_run = paragraph.add_run(base)
    base_run.italic = True
    _set_word_run_font(base_run, math=True)
    if index:
        index_run = paragraph.add_run(index)
        index_run.font.subscript = True
        _set_word_run_font(index_run, math=True)


def add_word_formula_legend(document: Document, trace: FormulaTrace) -> None:
    """Добавляет расшифровку всех обозначений непосредственно под формулой."""

    rows = formula_variable_rows(trace.formula_id)
    if not rows:
        return
    label = document.add_paragraph()
    label.paragraph_format.space_before = Pt(0)
    label.paragraph_format.space_after = Pt(1)
    label.paragraph_format.keep_with_next = True
    label_run = label.add_run("Обозначения:")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor.from_string("5B6770")

    definitions = document.add_paragraph()
    definitions.paragraph_format.space_before = Pt(0)
    definitions.paragraph_format.space_after = Pt(5)
    definitions.paragraph_format.keep_with_next = True
    for index, (symbol, title, unit) in enumerate(rows):
        _append_word_formula_symbol(definitions, symbol)
        suffix = f" — {title}"
        if unit and unit != "—":
            suffix += f", {unit}"
        suffix += "." if index == len(rows) - 1 else ";"
        text_run = definitions.add_run(suffix)
        _set_word_run_font(text_run)
        if index < len(rows) - 1:
            text_run.add_break()
    for run in definitions.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("374151")


def add_word_equation(
    document: Document,
    spec: MathSpec,
    *,
    keep_with_next: bool = False,
) -> None:
    """Добавляет центрированное редактируемое уравнение Office Math."""

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = True
    math = OxmlElement("m:oMath")
    _append_math(math, spec)
    paragraph._p.append(math)


def add_word_formula_block(
    document: Document,
    trace: FormulaTrace,
    *,
    heading_level: int,
) -> None:
    """Добавляет заголовок, формулу, подстановку, результат и источник."""

    heading = document.add_heading(trace.title_ru, level=heading_level)
    heading.paragraph_format.keep_with_next = True

    label = document.add_paragraph()
    label.paragraph_format.space_after = Pt(1)
    label.paragraph_format.keep_with_next = True
    label_run = label.add_run("Расчётная зависимость")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor.from_string("5B6770")

    spec = FORMULA_SPECS.get(trace.formula_id)
    if spec is None:
        spec = _run(_pretty_linear_text(trace.expression), italic=False)
    add_word_equation(document, spec, keep_with_next=True)
    add_word_formula_legend(document, trace)

    substitution_label = document.add_paragraph()
    substitution_label.paragraph_format.space_after = Pt(1)
    substitution_label.paragraph_format.keep_with_next = True
    substitution_run = substitution_label.add_run("Подстановка значений")
    substitution_run.bold = True
    substitution_run.font.size = Pt(9)
    substitution_run.font.color.rgb = RGBColor.from_string("5B6770")
    add_word_equation(
        document,
        _linear_math_spec(formatted_substitution_text(trace)),
        keep_with_next=True,
    )

    result_paragraph = document.add_paragraph()
    result_paragraph.paragraph_format.space_after = Pt(2)
    result_label = result_paragraph.add_run("Результат: ")
    result_label.bold = True
    result_paragraph.add_run(f"{trace.result:.2f} {trace.unit}")

    source_parts = [trace.standard]
    if trace.clause:
        source_parts.append(trace.clause)
    source = document.add_paragraph("Источник: " + ", ".join(source_parts) + ".")
    source.paragraph_format.space_after = Pt(9)
    for run in source.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("6B7280")


def pdf_formula_markup(trace: FormulaTrace) -> str:
    """Возвращает математически оформленную запись для ReportLab Paragraph."""

    return PDF_FORMULAS.get(
        trace.formula_id,
        _pretty_linear_text(trace.expression),
    )


def pdf_substitution_text(trace: FormulaTrace) -> str:
    """Возвращает компактную подстановку значений без программной нотации."""

    return formatted_substitution_text(trace)


def _pdf_formula_symbol_markup(symbol: str) -> str:
    prefix, base, index = _formula_symbol_parts(symbol)
    index_markup = f"<sub>{escape(index)}</sub>" if index else ""
    return f"{escape(prefix)}<i>{escape(base)}</i>{index_markup}"


def pdf_substitution_markup(trace: FormulaTrace) -> str:
    """Возвращает подстановку с настоящими подстрочными индексами для PDF."""

    parts: list[str] = []
    for kind, value in _linear_math_tokens(formatted_substitution_text(trace)):
        if kind == "symbol":
            parts.append(_pdf_formula_symbol_markup(value))
        elif kind == "round":
            parts.append("окр<sub>0,5↑</sub>")
        else:
            parts.append(escape(value))
    return "".join(parts)


def pdf_formula_legend_markup(trace: FormulaTrace) -> str:
    """Возвращает компактную расшифровку всех обозначений формулы для PDF."""

    rows = formula_variable_rows(trace.formula_id)
    if not rows:
        return ""
    definitions = []
    for symbol, title, unit in rows:
        suffix = f" — {escape(title)}"
        if unit and unit != "—":
            suffix += f", {escape(unit)}"
        definitions.append(f"{_pdf_formula_symbol_markup(symbol)}{suffix}")
    return "<b>Обозначения:</b><br/>" + ";<br/>".join(definitions) + "."


def formula_ids_with_native_layout() -> Sequence[str]:
    """Список формул, для которых подготовлена структурная разметка Office Math."""

    return tuple(FORMULA_SPECS)
