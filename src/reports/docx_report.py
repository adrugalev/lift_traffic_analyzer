"""Генератор подробного инженерного DOCX-отчёта."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src import APP_NAME, __version__
from src.models.project import Project
from src.models.results import CalculationResult, VariantResult
from src.models.simulation import SimulationResult
from src.reports.formula_rendering import add_word_formula_block


ACCENT = "2E74B5"
LIGHT_ACCENT = "F2F4F7"
TEXT = "1F2937"
DOCUMENT_LANGUAGE = "ru-RU"
COMPLIES_COLOR = "008000"
DOES_NOT_COMPLY_COLOR = "C00000"


def _set_cell_shading(cell: object, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_margins(cell: object) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    values = {"top": 80, "start": 120, "bottom": 80, "end": 120}
    for edge, value_dxa in values.items():
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value_dxa))
        node.set(qn("w:type"), "dxa")


def _style_compliance_cell(cell: object, value: object) -> None:
    """Выделяет итог нормативной оценки цветом и жирным шрифтом."""

    status = str(value).strip()
    color = {
        "Соответствует": COMPLIES_COLOR,
        "Не соответствует": DOES_NOT_COMPLY_COLOR,
    }.get(status)
    if color is None:
        return
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(color)


def _set_language_on_properties(properties: object) -> None:
    language = properties.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        properties.append(language)
    for attribute in ("val", "eastAsia", "bidi"):
        language.set(qn(f"w:{attribute}"), DOCUMENT_LANGUAGE)


def _set_document_language(document: Document) -> None:
    """Назначает русский язык всем стилям и текстовым фрагментам DOCX."""

    styles_element = document.styles.element
    defaults = styles_element.find(qn("w:docDefaults"))
    if defaults is None:
        defaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, defaults)
    run_defaults = defaults.find(qn("w:rPrDefault"))
    if run_defaults is None:
        run_defaults = OxmlElement("w:rPrDefault")
        defaults.insert(0, run_defaults)
    default_properties = run_defaults.find(qn("w:rPr"))
    if default_properties is None:
        default_properties = OxmlElement("w:rPr")
        run_defaults.append(default_properties)
    _set_language_on_properties(default_properties)

    parts = {document.part}
    for section in document.sections:
        parts.add(section.header.part)
        parts.add(section.footer.part)
    for part in parts:
        for run in part.element.iter(qn("w:r")):
            properties = run.find(qn("w:rPr"))
            if properties is None:
                properties = OxmlElement("w:rPr")
                run.insert(0, properties)
            _set_language_on_properties(properties)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal_fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal_fonts.set(qn(f"w:{attribute}"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Title", 23, "000000", 0, 4),
        ("Subtitle", 14, "373737", 0, 16),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            style_fonts.set(qn(f"w:{attribute}"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        if style_name == "Title":
            style_ppr = style._element.get_or_add_pPr()
            borders = style_ppr.find(qn("w:pBdr"))
            if borders is not None:
                style_ppr.remove(borders)
    _set_document_language(document)


def _set_table_geometry(table: object, widths: list[float]) -> None:
    """Фиксирует tblW/tblInd/tblGrid/tcW в DXA для ширины 9360."""

    total = sum(widths)
    dxa = [round(9360 * width / total) for width in widths]
    dxa[-1] += 9360 - sum(dxa)
    table_pr = table._tbl.tblPr
    for tag, attributes in (
        ("w:tblW", {"w:w": "9360", "w:type": "dxa"}),
        ("w:tblInd", {"w:w": "120", "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        node = table_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            table_pr.append(node)
        for key, value in attributes.items():
            node.set(qn(key), value)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)
    for row in table.rows:
        for cell, width in zip(row.cells, dxa, strict=True):
            cell_pr = cell._tc.get_or_add_tcPr()
            cell_width = cell_pr.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_pr.append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def _add_table(
    document: Document,
    headers: list[str],
    rows: list[list[object]],
    widths: list[float] | None = None,
    *,
    show_header: bool = True,
    add_spacing_after: bool = True,
    font_size_pt: float = 8.5,
) -> None:
    table = document.add_table(rows=1 if show_header else 0, cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"
    if show_header:
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = header
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_shading(cell, LIGHT_ACCENT)
            _set_cell_margins(cell)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor.from_string(TEXT)
                run.font.bold = True
                run.font.size = Pt(font_size_pt)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cells[index])
            if row_index % 2:
                _set_cell_shading(cells[index], "F4F7F9")
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size_pt)
            _style_compliance_cell(cells[index], value)
    _set_table_geometry(table, widths or [1.0] * len(headers))
    if add_spacing_after:
        document.add_paragraph()


def _add_list_item(document: Document, text: str, numbered: bool = False) -> None:
    """Добавляет настоящий Word-список с токенами standard_business_brief."""

    paragraph = document.add_paragraph(
        text,
        style="List Number" if numbered else "List Bullet",
    )
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def build_docx_report(
    project: Project,
    analytic: CalculationResult | None = None,
    simulation: SimulationResult | None = None,
    variants: list[VariantResult] | None = None,
) -> bytes:
    """Формирует самостоятельный DOCX-отчёт и возвращает его содержимое."""

    document = Document()
    _configure_document(document)
    header = document.sections[0].header.paragraphs[0]
    header.text = "ИНЖЕНЕРНЫЙ ОТЧЁТ  |  АНАЛИЗ ПАССАЖИРОПОТОКА"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("6B7280")
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Анализ пассажиропотока\nлифтовой группы")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run(project.metadata.name)
    document.add_paragraph(f"Объект: {project.metadata.address or 'не указан'}")
    document.add_paragraph(f"Заказчик: {project.metadata.customer or 'не указан'}")
    document.add_paragraph(f"Проектировщик: {project.metadata.designer or 'не указан'}")
    document.add_paragraph(f"Разработчик отчёта: {project.metadata.calculation_author or 'не указан'}")
    document.add_paragraph(f"Дата расчёта: {project.metadata.calculation_date.isoformat()}")
    is_gost = bool(
        analytic and analytic.calculation_basis == "GOST_34758_2021_CLAUSE_7"
    )
    warning = document.add_paragraph()
    warning.style = document.styles["Normal"]
    if is_gost:
        failed = any(
            metric.compliance.value == "Не соответствует"
            for metric in analytic.metrics
        )
        status_text = (
            "Статус: выполнен расчётный метод ГОСТ 34758-2021; "
            f"результат — {'имеются несоответствия' if failed else 'критерии выполнены'}."
        )
    else:
        status_text = (
            "Статус: нормативное соответствие не оценено, пока не выполнен "
            "расчёт по верифицированной нормативной конфигурации."
        )
    run = warning.add_run(status_text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("9C5700")

    document.add_page_break()
    document.add_heading("Содержание", level=1)
    sections = [
        "Executive Summary",
        "Сведения об объекте",
        "Принятые исходные данные",
        "Распределение населения",
        "Описание лифтовых групп",
        "Расчётный пассажиропоток",
        "Нормативные критерии",
        "Аналитический расчёт",
        "Формулы и подстановка значений",
        "Результаты симуляции",
        "Сравнение вариантов",
        "Выводы и рекомендации",
        "Ограничения расчёта",
        "Версия программы и аудит",
    ]
    for name in sections:
        _add_list_item(document, name, numbered=True)

    document.add_heading("Executive Summary", level=1)
    if analytic:
        document.add_paragraph(
            (
                "Выполнен расчётный метод ГОСТ 34758-2021 для пикового "
                "пассажиропотока вверх. Нормативная оценка показателей приведена ниже."
            )
            if is_gost
            else (
                "Выполнен предварительный инженерный расчёт. Полученные показатели "
                "служат для сравнения вариантов и постановки симуляции, но не являются "
                "заключением о соответствии ГОСТ 34758-2021."
            )
        )
        _add_table(
            document,
            ["Показатель", "Значение", "Метод", "Статус"],
            [
                [
                    metric.title_ru,
                    f"{metric.value:.2f} {metric.unit}",
                    metric.method,
                    metric.compliance.value,
                ]
                for metric in analytic.metrics
                if metric.key in {"interval", "handling_capacity_5min", "average_wait_proxy", "reserve"}
            ],
            [2.6, 1.3, 1.9, 1.1],
        )
    else:
        document.add_paragraph("Аналитический расчёт не приложен к текущему отчёту.")

    document.add_heading("Сведения об объекте", level=1)
    _add_table(
        document,
        ["Поле", "Значение"],
        [
            ["Название", project.metadata.name],
            ["Адрес", project.metadata.address or "не указан"],
            ["Тип здания", project.building.building_type.value],
            ["Стадия", project.metadata.design_stage],
            ["Выбранный стандарт", project.metadata.selected_standard.value],
            ["Система единиц", project.metadata.units.value],
        ],
        [2.2, 4.7],
    )

    document.add_heading("Принятые исходные данные", level=1)
    document.add_paragraph(
        "Значения технических параметров введены пользователем или взяты из редактируемого "
        "стартового шаблона. Они не являются данными конкретного производителя."
    )
    document.add_heading("Распределение населения", level=2)
    document.add_paragraph(
        f"Население при 100%: {project.base_population} чел. "
        f"Коэффициент заселённости: {project.building.occupancy_percent}%. "
        f"Расчётное население: {project.population} чел."
    )
    _add_table(
        document,
        ["Этаж", "Отметка, м", "Назначение", "Расч. население", "Паркинг"],
        [
            [
                floor.label or floor.number,
                f"{floor.elevation_m:.2f}",
                floor.purpose,
                f"{project.effective_floor_population(floor):.1f}",
                "Да" if floor.is_parking else "Нет",
            ]
            for floor in project.floors
        ],
        [0.8, 1.2, 2.8, 1.2, 1.0],
    )

    document.add_heading("Описание лифтовых групп", level=1)
    for group in project.elevator_groups:
        document.add_heading(group.name, level=2)
        document.add_paragraph(
            f"Основной этаж: {group.main_floor}. "
            f"Обслуживаемые этажи: {', '.join(map(str, group.served_floors))}."
        )
        _add_table(
            document,
            ["Лифт", "Г/п, кг", "Пасс.", "Скорость, м/с", "Двери, м", "Заполнение"],
            [
                [
                    elevator.name,
                    f"{elevator.capacity_kg:.0f}",
                    elevator.nominal_passengers,
                    f"{elevator.speed_mps:.2f}",
                    f"{elevator.door_width_m:.2f}",
                    f"{elevator.load_factor:.0%}",
                ]
                for elevator in group.elevators
            ],
            [1.5, 0.9, 0.7, 1.1, 0.9, 1.0],
        )
        _add_table(
            document,
            ["Лифт", "Ускорение, м/с²", "Замедление, м/с²", "Рывок, м/с³"],
            [
                [
                    elevator.name,
                    f"{elevator.acceleration_mps2:.2f}",
                    f"{elevator.deceleration_mps2:.2f}",
                    f"{elevator.jerk_mps3:.2f}",
                ]
                for elevator in group.elevators
            ],
            [1.8, 1.5, 1.5, 1.2],
        )

    scenario = project.scenario()
    parking_description = (
        f" Доля входящего потока с паркинга: {scenario.parking_incoming_share:.0%}."
        if any(floor.is_parking for floor in project.floors)
        else ""
    )
    document.add_heading("Расчётный пассажиропоток", level=1)
    document.add_paragraph(
        f"Сценарий: {scenario.scenario_type.value}. "
        f"Пятиминутный процент: {scenario.population_percent_5min:.2f} %. "
        f"Доли входящего/исходящего/межэтажного потока: "
        f"{scenario.incoming_share:.0%}/{scenario.outgoing_share:.0%}/{scenario.interfloor_share:.0%}. "
        f"{parking_description}"
    )

    document.add_heading("Нормативные критерии", level=1)
    if is_gost:
        document.add_paragraph(
            "Критерии приняты по таблицам 1 и 4 ГОСТ 34758-2021 для выбранного "
            "назначения здания."
        )
        _add_table(
            document,
            ["Показатель", "Факт", "Требование", "Оценка"],
            [
                [
                    metric.title_ru,
                    f"{metric.value:.2f} {metric.unit}",
                    metric.target_description or "—",
                    metric.compliance.value,
                ]
                for metric in analytic.metrics
                if metric.key
                in {
                    "handling_capacity_5min",
                    "specific_capacity",
                    "interval",
                    "full_height_time",
                }
            ],
            [2.5, 1.5, 1.5, 1.4],
        )
    else:
        document.add_paragraph(
            "Критерии не приведены: нормативный расчёт не приложен. "
            "В отчёте намеренно отсутствуют придуманные пороги и номера пунктов."
        )

    document.add_heading("Аналитический расчёт", level=1)
    if analytic:
        _add_table(
            document,
            ["Показатель", "Результат", "Единица", "Нормативная оценка"],
            [
                [metric.title_ru, f"{metric.value:.3f}", metric.unit, metric.compliance.value]
                for metric in analytic.metrics
            ],
            [3.0, 1.2, 1.3, 1.4],
        )
        document.add_heading("Формулы и подстановка значений", level=2)
        for trace in analytic.formulas:
            add_word_formula_block(document, trace, heading_level=3)
    else:
        document.add_paragraph("Расчёт не выполнен.")

    document.add_heading("Результаты симуляции", level=1)
    if simulation:
        _add_table(
            document,
            ["Показатель", "Значение"],
            [
                ["Среднее время ожидания пассажира (AWT)", f"{simulation.waiting_time.mean:.2f} с"],
                ["Медианное время ожидания (P50)", f"{simulation.waiting_time.median:.2f} с"],
                [
                    "Ожидание 95% пассажиров, не более (P95)",
                    f"{simulation.waiting_time.percentile_95:.2f} с",
                ],
                [
                    "Среднее полное время до этажа назначения (TTD)",
                    f"{simulation.time_to_destination.mean:.2f} с",
                ],
                ["Средняя очередь, пассажиров", f"{simulation.average_queue_length:.2f} пасс."],
                ["Максимальная очередь, пассажиров", simulation.maximum_queue_length],
                ["Перевезено пассажиров (среднее за повтор)", simulation.transported_passengers],
                ["Не обслужено пассажиров (среднее за повтор)", simulation.unserved_passengers],
            ],
            [4.3, 2.4],
        )
    else:
        document.add_paragraph("Симуляция не выполнялась или не приложена.")

    document.add_heading("Сравнение вариантов", level=1)
    if variants:
        _add_table(
            document,
            ["Вариант", "Лифты", "Г/п", "Скорость", "Интервал", "HC5", "Оценка"],
            [
                [
                    item.variant_name,
                    item.elevator_count,
                    f"{item.capacity_kg:.0f}",
                    f"{item.speed_mps:.1f}",
                    f"{item.interval_s:.1f}",
                    f"{item.handling_capacity_5min:.1f}",
                    f"{item.score:.1f}",
                ]
                for item in variants
            ],
            [1.4, 0.6, 0.7, 0.9, 0.9, 0.9, 0.9],
        )
    else:
        document.add_paragraph("Сохранённые варианты отсутствуют.")

    document.add_heading("Выводы и рекомендации", level=1)
    if analytic and analytic.recommendations:
        for recommendation in analytic.recommendations:
            _add_list_item(
                document,
                f"{recommendation.problem} Предлагаемое действие: {recommendation.proposed_action} "
                f"Ожидаемый эффект: {recommendation.expected_effect}",
            )
    else:
        document.add_paragraph("Расчётные рекомендации отсутствуют.")

    document.add_heading("Ограничения расчёта", level=1)
    if analytic and analytic.calculation_basis == "GOST_34758_2021_CLAUSE_7":
        limitations = (
            "Расчётный метод применим к восходящему пику, одному нижнему входу, "
            "равномерному заселению и однородной группе.",
            "Формулы ГОСТ сверены с открытой публикацией; договорный отчёт требует контроля "
            "по экземпляру стандарта заказчика.",
            "Ориентировочное время ожидания не является нормативным критерием ГОСТ.",
            "Симуляция MVP использует прозрачную упрощённую диспетчеризацию.",
            "Модель энергопотребления отсутствует.",
        )
    else:
        limitations = (
            "Предварительный аналитический цикл не является нормативным RTT.",
            "Симуляция MVP использует прозрачную упрощённую диспетчеризацию.",
            "Модель энергопотребления отсутствует.",
        )
    for item in limitations:
        _add_list_item(document, item)

    document.add_heading("Версия программы и аудит", level=1)
    document.add_paragraph(f"Версия приложения: {__version__}")
    if analytic:
        document.add_paragraph(f"Хэш проекта: {analytic.audit.project_hash}")
        document.add_paragraph(f"Версия конфигурации: {analytic.audit.configuration_version}")
    if simulation:
        document.add_paragraph(
            f"Симуляция: seed={simulation.seed}, повторов={simulation.repetitions}, "
            f"хэш проекта={simulation.project_hash}"
        )

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"{APP_NAME} {__version__} • локальный расчёт")
    _set_document_language(document)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()
