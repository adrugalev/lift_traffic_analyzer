"""Интеграционные тесты полного пути проекта и экспортов."""

from __future__ import annotations

from io import BytesIO
import re
from xml.etree import ElementTree
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook
from pypdf import PdfReader

from src.engines.analytic_engine import AnalyticEngine
from src.engines.simulation_engine import SimulationEngine
from src.models.simulation import SimulationSettings
from src.reports.docx_report import build_docx_report
from src.reports.excel_report import build_excel_report
from src.reports.formula_rendering import formula_ids_with_native_layout
from src.reports.gost_report import build_gost_docx_report, build_gost_pdf_report
from src.reports.pdf_report import build_pdf_report
from src.services.project_service import ProjectService


MATH_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _word_math_elements(data: bytes, name: str) -> list[ElementTree.Element]:
    with ZipFile(BytesIO(data)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    return root.findall(f".//{{{MATH_NAMESPACE}}}{name}")


def _word_equation_texts(data: bytes) -> list[str]:
    return [
        "".join(node.text or "" for node in equation.iter())
        for equation in _word_math_elements(data, "oMath")
    ]


def _assert_centered_word_equations(data: bytes, expected_count: int) -> None:
    with ZipFile(BytesIO(data)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    equation_paragraphs = [
        paragraph
        for paragraph in root.findall(f".//{{{WORD_NAMESPACE}}}p")
        if paragraph.find(f".//{{{MATH_NAMESPACE}}}oMath") is not None
    ]
    assert len(equation_paragraphs) == expected_count
    for paragraph in equation_paragraphs:
        alignment = paragraph.find(
            f"./{{{WORD_NAMESPACE}}}pPr/{{{WORD_NAMESPACE}}}jc"
        )
        assert alignment is not None
        assert alignment.get(f"{{{WORD_NAMESPACE}}}val") == "center"


def _assert_russian_word_language(data: bytes) -> None:
    with ZipFile(BytesIO(data)) as package:
        document_root = ElementTree.fromstring(package.read("word/document.xml"))
        styles_root = ElementTree.fromstring(package.read("word/styles.xml"))
    language_tag = f"{{{WORD_NAMESPACE}}}lang"
    run_properties_tag = f"{{{WORD_NAMESPACE}}}rPr"
    run_tag = f"{{{WORD_NAMESPACE}}}r"
    value_attribute = f"{{{WORD_NAMESPACE}}}val"
    east_asia_attribute = f"{{{WORD_NAMESPACE}}}eastAsia"
    bidi_attribute = f"{{{WORD_NAMESPACE}}}bidi"

    defaults_language = styles_root.find(
        f".//{{{WORD_NAMESPACE}}}docDefaults/"
        f"{{{WORD_NAMESPACE}}}rPrDefault/"
        f"{{{WORD_NAMESPACE}}}rPr/{language_tag}"
    )
    assert defaults_language is not None
    assert defaults_language.get(value_attribute) == "ru-RU"
    assert defaults_language.get(east_asia_attribute) == "ru-RU"
    assert defaults_language.get(bidi_attribute) == "ru-RU"

    for run in document_root.iter(run_tag):
        properties = run.find(run_properties_tag)
        assert properties is not None
        language = properties.find(language_tag)
        assert language is not None
        assert language.get(value_attribute) == "ru-RU"


def test_project_json_round_trip_has_no_data_loss() -> None:
    project = ProjectService.create_default()
    restored = ProjectService.loads(ProjectService.dump_bytes(project))
    assert restored.model_dump(exclude={"modified_at"}) == project.model_dump(exclude={"modified_at"})


def test_docx_export_opens_and_contains_method_boundary() -> None:
    project = ProjectService.create_default()
    analytic = AnalyticEngine().calculate_preview(project)
    data = build_docx_report(project, analytic)
    document = Document(BytesIO(data))
    text = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )
    assert "нормативное соответствие не оценено" in text.lower()
    assert "Формулы и подстановка значений" in text
    assert "Ускорение, м/с²" in text
    assert "Замедление, м/с²" in text
    assert "Рывок, м/с³" in text
    assert len(_word_math_elements(data, "oMath")) == 2 * len(analytic.formulas)
    _assert_centered_word_equations(data, 2 * len(analytic.formulas))
    assert _word_math_elements(data, "f")
    assert _word_math_elements(data, "nary")
    _assert_russian_word_language(data)
    assert {trace.formula_id for trace in analytic.formulas} <= set(
        formula_ids_with_native_layout()
    )
    with ZipFile(BytesIO(data)) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
    assert "окр" in document_xml
    assert "round" not in document_xml
    assert text.count("Обозначения:") >= len(analytic.formulas)
    for trace in analytic.formulas:
        assert f"{trace.result:.2f} {trace.unit}" in text
    substitution_texts = _word_equation_texts(data)[1::2]
    assert len(substitution_texts) == len(analytic.formulas)
    assert not any(re.search(r"\d+\.\d{3,}", item) for item in substitution_texts)


def test_xlsx_export_has_required_sheets_and_typed_values() -> None:
    project = ProjectService.create_default()
    analytic = AnalyticEngine().calculate_preview(project)
    simulation = SimulationEngine().run(project, SimulationSettings(repetitions=1, random_seed=1))
    workbook = load_workbook(BytesIO(build_excel_report(project, analytic, simulation)))
    assert workbook.sheetnames == [
        "Исходные данные",
        "Этажи",
        "Лифты",
        "Формулы",
        "Аналитический расчёт",
        "Симуляция",
        "Сравнение",
        "Рекомендации",
    ]
    assert isinstance(workbook["Этажи"]["F2"].value, int)
    assert workbook["Лифты"]["I1"].value == "Рывок, м/с³"
    assert isinstance(workbook["Аналитический расчёт"]["C2"].value, (int, float))


def test_pdf_export_reopens_and_has_pages() -> None:
    project = ProjectService.create_default()
    analytic = AnalyticEngine().calculate_preview(project)
    data = build_pdf_report(project, analytic)
    reader = PdfReader(BytesIO(data))
    assert len(reader.pages) >= 2
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "Ускорение, м/с²" in text
    assert "Замедление, м/с²" in text
    assert "Рывок, м/с³" in text


def test_gost_report_contains_section_9_information() -> None:
    project = ProjectService.create_default()
    analytic = AnalyticEngine().calculate_normative(project)
    data = build_gost_docx_report(project, analytic)
    document = Document(BytesIO(data))
    text_parts = [paragraph.text for paragraph in document.paragraphs]
    text_parts.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    text = "\n".join(text_parts)
    assert "Форма по ГОСТ 34758-2021, раздел 9" in text
    assert "1. Общие сведения" in text
    assert "2. Информация о здании" in text
    assert "3. Критерии проектирования" in text
    assert "4. Информация о лифтовой установке" in text
    assert "5. Расчётные данные провозной способности" in text
    assert "Итоговая оценка соответствия" in text
    assert "Время предварительного открывания" in text
    assert "Тип открывания дверей" in text
    assert "Ускорение" in text
    assert "Замедление" in text
    assert "Рывок" in text
    assert "Достижение номинальной скорости на межэтажном пролёте" in text
    assert "Максимальная скорость на межэтажном пролёте" in text
    assert "Проверка достижения номинальной скорости" in text
    assert "Межэтажное время движения с разгоном и торможением" in text
    assert "0,5 округляется вверх" in text
    assert document.tables[0].cell(0, 0).text == "Компания, выполнившая проектирование"
    assert document.tables[0].cell(0, 1).text == (project.metadata.designer or "не указана")
    assert document.tables[1].cell(0, 0).text == "Наименование здания"
    assert document.tables[1].cell(0, 1).text == project.metadata.name
    assert "Сведения по п. 9" not in text
    assert "Критерий по п. 9" not in text
    report_datetime = analytic.audit.calculated_at.astimezone().strftime(
        "%d.%m.%Y %H:%M:%S"
    )
    assert report_datetime in text
    assert f"Метод: {analytic.method}." in text_parts
    assert f"Дата и время расчёта: {report_datetime}." in text_parts
    assert not any(
        paragraph.startswith("Метод:") and "Дата и время расчёта:" in paragraph
        for paragraph in text_parts
    )
    assert "RTZ" not in text
    assert "(зима)" not in text
    assert len(_word_math_elements(data, "oMath")) == 2 * len(analytic.formulas)
    _assert_centered_word_equations(data, 2 * len(analytic.formulas))
    assert _word_math_elements(data, "f")
    assert _word_math_elements(data, "nary")
    _assert_russian_word_language(data)
    assert {trace.formula_id for trace in analytic.formulas} <= set(
        formula_ids_with_native_layout()
    )
    assert text.count("Обозначения:") >= len(analytic.formulas)
    for trace in analytic.formulas:
        assert f"{trace.result:.2f} {trace.unit}" in text
    substitution_texts = _word_equation_texts(data)[1::2]
    assert len(substitution_texts) == len(analytic.formulas)
    assert not any(re.search(r"\d+\.\d{3,}", item) for item in substitution_texts)
    with ZipFile(BytesIO(data)) as package:
        document_root = ElementTree.fromstring(package.read("word/document.xml"))
    subscript_attribute = f"{{{WORD_NAMESPACE}}}val"
    assert any(
        element.get(subscript_attribute) == "subscript"
        for element in document_root.findall(
            f".//{{{WORD_NAMESPACE}}}vertAlign"
        )
    )

    calculation_table = next(
        table
        for table in document.tables
        if table.cell(1, 0).text == "Этаж реверса"
    )
    assert [cell.text for cell in calculation_table.rows[0].cells] == [
        "Параметр",
        "Результат",
        "Критерий",
        "Оценка",
    ]

    floors_table = next(
        table
        for table in document.tables
        if table.cell(0, 0).text == "Этаж"
    )
    floor_widths = [
        int(column.get(qn("w:w")))
        for column in floors_table._tbl.tblGrid
    ]
    assert floor_widths[1] >= 1200
    assert floor_widths[5] >= 1500

    criteria_table = next(
        table
        for table in document.tables
        if len(table.rows) > 1
        and table.cell(1, 0).text == "Провозная способность для пикового входящего потока"
    )
    assert criteria_table.cell(0, 0).text == "Параметр"

    installation_table = next(
        table
        for table in document.tables
        if len(table.rows) > 1 and table.cell(1, 0).text == "Расположение группы"
    )
    assert installation_table.cell(0, 0).text == "Параметр"


def test_gost_pdf_report_reopens_and_has_pages() -> None:
    project = ProjectService.create_default()
    analytic = AnalyticEngine().calculate_normative(project)
    data = build_gost_pdf_report(project, analytic)
    reader = PdfReader(BytesIO(data))
    assert len(reader.pages) >= 2
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "Сведения по п. 9" not in text
    assert "Критерий по п. 9" not in text
    report_datetime = analytic.audit.calculated_at.astimezone().strftime(
        "%d.%m.%Y %H:%M:%S"
    )
    assert report_datetime in text
    assert "окр" in text
    assert "round" not in text
    assert "Максимальная скорость на межэтажном пролёте" in text
    assert "Проверка достижения номинальной скорости" in text
    assert text.count("Обозначения:") >= len(analytic.formulas)
    for trace in analytic.formulas:
        assert f"{trace.result:.2f} {trace.unit}" in text
    assert not re.search(r"Результат:\s*-?\d+\.\d{3,}", text)
    assert "RTZ" not in text
    assert "(зима)" not in text
