"""Проверки цветового оформления нормативных статусов в DOCX."""

from __future__ import annotations

from docx import Document

from src.reports.docx_report import (
    COMPLIES_COLOR,
    DOES_NOT_COMPLY_COLOR,
    _add_table,
)


def test_compliance_statuses_are_bold_and_colored() -> None:
    """Положительный и отрицательный статусы визуально различимы."""

    document = Document()
    _add_table(
        document,
        ["Параметр", "Статус"],
        [
            ["Интервал", "Соответствует"],
            ["Провозная способность", "Не соответствует"],
        ],
    )

    table = document.tables[0]
    expected = {
        "Соответствует": COMPLIES_COLOR,
        "Не соответствует": DOES_NOT_COMPLY_COLOR,
    }
    for row in table.rows[1:]:
        cell = row.cells[1]
        run = cell.paragraphs[0].runs[0]
        assert run.bold is True
        assert str(run.font.color.rgb) == expected[cell.text]
